import logging
import pathlib
import uuid
from typing import List

import docx
from fastapi import (
    APIRouter,
    BackgroundTasks,
    HTTPException,
    Request,
    UploadFile,
)
from fastapi.responses import FileResponse

# from pptx_helper import generate_powerpoint_presentation
from pydantic import BaseModel

# TODO: only here for testing purposes
from model.chunk_and_vectorise import chunk_and_vectorise
from model.dao import (
    get_all_chunks,
    get_all_topics,
    get_name_from_datbase,
    remove_by_learning_id,
    remove_document_from_source_documents,
    send_chunks_to_db,
    send_learning_to_db,
)
from model.gail_util_funcs import (
    delete_temp_file,
    list_to_markdown,
    markdown_to_list,
    transform_string_of_bullets_to_list,
)
from model.generate_powerpoint.generate_powerpoint_presentation import (
    generate_powerpoint_presentation,
)
from model.get_relevant_chunks import get_relevant_chunks
from model.pii_mask import guardrail_check
from model.process_uploaded_file import extract_markdown
from model.scan_file_for_malware import scan_file_for_malware

logger = logging.getLogger(__name__)

doc_manager_router_v1 = APIRouter()

NO_LEARNING_ID = "No-learning-id"
WORD_DOC_TEMPLATE_PATH = "src/data/docx_templates/DWP_Template.docx"
POWERPOINT_TEMPLATE_PATH = "src/data/pptx_templates/DWP_Template.pptx"


class QueryString(BaseModel):
    query_strings: List[str]
    top_k: int


class NamePurposeTopics(BaseModel):
    name: str
    purpose: str
    topics: List[str]


class vectorisedChunkObject(BaseModel):
    doc_id: uuid.UUID
    doc_title: str
    chunk_id: int
    raw_text: str
    vector: List[float]


class FacilitatorSlide(BaseModel):
    script: str
    bullet_points: str
    heading: str


class KnowledgeCheck(BaseModel):
    question: str
    choices: str
    answer: str


class TopicOutlines(BaseModel):
    title: str
    objectives: str


class DocumentSlides(BaseModel):
    slides: List[FacilitatorSlide]


class DocumentQuestions(BaseModel):
    questions: List[KnowledgeCheck]


class DocumentContent(BaseModel):
    topic_outline: TopicOutlines
    content: DocumentQuestions


class PowerPointContent(BaseModel):
    topic_outline: TopicOutlines
    content: DocumentSlides


class DownloadDocx(BaseModel):
    extension: str
    module_title: str
    topics: List[DocumentContent]


class DownloadPptx(BaseModel):
    extension: str
    module_title: str
    topics: List[PowerPointContent]


def log_incoming_request(request: Request):
    reqid = request.headers.get("X-Request-ID")
    reqorigin = request.headers.get("X-Origin")

    logger.info("Recieved request_id: %s \n with Origin: %s", reqid, reqorigin)


@doc_manager_router_v1.post("/uploadnameandpurpose")
async def upload_name_purpose(request: Request, data: NamePurposeTopics):
    log_incoming_request(request)
    logger.info("Entered upload_name_purpose")
    learning_id = request.headers.get("Learning-ID")

    if learning_id == NO_LEARNING_ID:
        learning_id = None

    try:
        return await send_learning_to_db(data, learning_id)

    except Exception as error:
        logger.error(f"Error occured in upload_name_purpose: {error}")
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred in upload_name_purpose: {error}",
        )


@doc_manager_router_v1.post("/scanfile")
async def scan_file(request: Request, file: UploadFile):
    log_incoming_request(request)
    logger.info("Entered scan_file")

    try:
        file_content = await file.read()
        # Intercept headers from request where required.
        response = await scan_file_for_malware(file, file_content)

        logger.info("Scan file response recieved", response)
        return response

    except HTTPException as http_error:
        logger.info(
            "Malware detected in file %s: %s",
            file.filename,
            http_error.detail,
        )
    except Exception:
        logger.error("Error in network")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while processing the file",
        )


@doc_manager_router_v1.post("/uploadfile")
async def upload_file(request: Request, file: UploadFile):
    log_incoming_request(request)
    learning_id = request.headers.get("Learning-ID")
    logger.info("Entered upload_file")
    # Intercept headers from request where required.

    try:
        logger.info("Successfully uploaded %s", file.filename)
        logger.info("Reading file %s", file.filename)
        readFile = await file.read()

        # Scan file for Malware and respond if found
        logger.info("Attempting to scan file %s", file.filename)
        scan_response = await scan_file_for_malware(file, readFile)
        infected = scan_response[0]["is_infected"]
        if infected:
            logger.info("Malware detected %s", file.filename)
            return {
                "pii_detected": False,
                "success": bool(scan_response),
                "is_infected": infected,
            }

        #  Extract Markdown
        markdown_text = await extract_markdown(file, readFile)
        # Guardrail check and remove any PII
        pii_detected, masked_md_text = guardrail_check(markdown_text, "OUTPUT")

        # chunk the document
        vectorised_chunks, doc_id = await chunk_and_vectorise(
            masked_md_text, doc_title=file.filename
        )
        response = await send_chunks_to_db(learning_id, vectorised_chunks)
        return {
            "pii_detected": pii_detected,
            "success": bool(response),
            "document_id": doc_id,
            "is_infected": infected,
        }

    except Exception as error:
        logger.error(f"Error occured in upload_file: {error}")
        return error


@doc_manager_router_v1.post("/getrelevantchunks")
async def fetch_relevant_chunks(request: Request, data: QueryString):
    log_incoming_request(request)
    learning_id = request.headers.get("Learning-ID")
    logger.info("Entered get relevant chunks from database")
    try:
        # Get all chunks from the database
        all_chunks = await get_all_chunks(learning_id)
        logger.info("All chunks recieved")
        logger.info("Query strings recieved")

        response = await get_relevant_chunks(
            all_chunks, data.query_strings, data.top_k
        )
        logger.info("Relevant chunks received")
        return {"relevant_chunks": response}

    except Exception as error:
        logger.error("Error in network")
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while getting the relevant chunks: {error}",
        )


@doc_manager_router_v1.get("/topicgenerationvariables")
async def get_variables_for_creation_ms(request: Request):
    log_incoming_request(request)
    learning_id = request.headers.get("Learning-ID")

    logger.info("Entered get variables from database for content creation MS")
    # TODO error handling with empty learningID
    try:
        query_strings = await get_all_topics(learning_id)
        logger.info("Query strings recieved")
        name = await get_name_from_datbase(learning_id)
        logger.info("Name of journey received")
        return {"name": name, "topic_outlines": query_strings}

    except Exception as error:
        logger.error("Error in network")
        raise HTTPException(
            status_code=500,
            detail=f"An error occurred while getting the relevant chunks: {error}",
        )


@doc_manager_router_v1.delete("/removedocument")
async def remove_document(request: Request):
    log_incoming_request(request)
    logger.info("Entered remove_document")
    learning_id = request.headers.get("Learning-ID")

    if not learning_id:
        logger.error(
            "Error occured in remove_document: missing learning_id header"
        )
        raise HTTPException(
            status_code=400, detail="Missing learning_id header"
        )

    try:
        result = await remove_by_learning_id(learning_id)
        return result
    except Exception as error:
        logger.error(f"Error occured in remove_document: {error}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while removing document",
        )


@doc_manager_router_v1.delete("/removesourcedocument")
async def remove_source_document(request: Request):
    log_incoming_request(request)
    logger.info("Entered remove_source_document")
    learning_id = request.headers.get("Learning-ID")
    document_id = request.headers.get("Document-ID")

    if not learning_id or not document_id:
        raise HTTPException(
            status_code=400,
            detail="Missing learning_id or document_id header",
        )

    try:
        result = await remove_document_from_source_documents(
            learning_id, document_id
        )
        return result

    except Exception as error:
        logger.error(f"Error occured in remove__source_document: {error}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while removing source document",
        )


@doc_manager_router_v1.get("/liveness/")
async def get_liveness(request: Request):
    return {"message": "TODO: implement liveness route"}


@doc_manager_router_v1.get("/readiness/")
async def get_readiness(request: Request):
    return {"message": "TODO: implement readiness route"}


@doc_manager_router_v1.post("/writeknowledgecheck")
async def write_kc_to_file(
    request: Request, data: DownloadDocx, background_tasks: BackgroundTasks
):
    logger.info("Entered write_kc_to_file")
    learning_id = request.headers.get("Learning-ID")

    if not learning_id:
        raise HTTPException(
            status_code=400,
            detail="Missing learning_id header",
        )

    try:
        json_body = data.model_dump()

        file_name = json_body["module_title"]
        extension = json_body["extension"]

        # Create new doc
        doc = docx.Document(WORD_DOC_TEMPLATE_PATH)

        # Set the header as the module title
        header = doc.sections[0].header
        header.add_paragraph(json_body["module_title"], style="Header White")

        # Add a document title
        doc.add_heading("Knowledge Checker:", 1)
        doc.add_heading(f"'{json_body['module_title']}'")
        doc.add_paragraph("")
        # Iterate through topics
        for i, kc in enumerate(json_body["topics"]):
            # Add topic heading
            doc.add_heading(
                f"Topic {str(i + 1)}: {kc['topic_outline']['title']}",
                level=2,
            )
            # Add questions
            for i, q in enumerate(kc["content"]["questions"]):
                doc.add_paragraph()  # Small break between Qs

                q_para = doc.add_paragraph(style="Body text 14pt")
                question_num = q_para.add_run(f"Question {str(i + 1)}. ")
                question_num.bold = True
                q_para.add_run(q["question"])

                if "choices" in q:
                    if isinstance(q["choices"], list):
                        q["choices"] = (
                            list_to_markdown(q["choices"])
                            .strip()
                            .replace("-", "")
                        )
                    else:
                        q["choices"] = q["choices"].strip().replace("-", "")
                    choices_para = doc.add_paragraph(style="Body text 14pt")
                    choices_title = choices_para.add_run("Choices: \n")
                    choices_title.bold = True
                    choices_para.add_run(q["choices"])

                a_para = doc.add_paragraph(style="Body text 14pt")
                if isinstance(q["answer"], list):  # For multi-answer questions
                    q["answer"] = (
                        list_to_markdown(q["answer"]).strip().replace("-", "")
                    )
                    answer_title = a_para.add_run(f"Answer: \n{q['answer']}")
                    answer_title.bold = True
                else:
                    answer_title = a_para.add_run(f"Answer: {q['answer']}")
                    answer_title.bold = True
            # Add page break
            doc.add_page_break()
        file_name = file_name.replace(" ", "_")
        # Create temp directory
        persistant_dir = pathlib.Path("/tmp/persitant_files") / learning_id
        persistant_dir.mkdir(parents=True, exist_ok=True)

        final_file_path = persistant_dir / f"{file_name}.{extension}"

        doc.save(final_file_path)

        background_tasks.add_task(delete_temp_file, final_file_path)

        # Send powepoint back to UI using FileResponse
        return FileResponse(
            pathlib.Path(final_file_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=file_name,
        )
    except Exception as error:
        logger.error(f"Error occured in write_kc_to_file: {error}")
        raise HTTPException(
            status_code=500,
            detail="An error occurred while create word document",
        )


@doc_manager_router_v1.post("/writepowerpoint")
async def write_facilitator_pptx_to_file(
    request: Request, data: DownloadPptx, background_tasks: BackgroundTasks
):
    logger.info("Entered write_facilitator_pptx_to_file")
    learning_id = request.headers.get("Learning-ID")

    if not learning_id:
        raise HTTPException(
            status_code=400,
            detail="Missing learning_id header",
        )
    try:
        json_body = data.model_dump()
        file_name = json_body["module_title"]
        extension = json_body["extension"]
        topics = json_body["topics"]

        export_pptx_json = {
            "title": file_name,
            "slides": [],
        }

        for i, topic in enumerate(topics):
            topic_num = i + 1

            topic_outline = topic["topic_outline"]
            topic_content = topic["content"]

            # Ensure correct formatting of bulletpoints (nested list)
            topic_content["slides"] = transform_string_of_bullets_to_list(
                topic_content["slides"]
            )

            # Append slides to entire pptx
            # Add new topic title slide
            topic_title = topic_outline["title"]
            if isinstance(topic_outline["objectives"], str):
                # Make sure objectives are a list, not a string! (else won't
                # properly format as bulletpoints on export)
                topic_outline["objectives"] = markdown_to_list(
                    topic_outline["objectives"]
                )
            topic_title_slide = {
                "topic_num": f"Topic {str(topic_num)}",
                "title": topic_title,
                "script": "",
            }
            export_pptx_json["slides"].append(topic_title_slide)
            topic_objectives_slide = {
                "heading": f"Topic {str(topic_num)}: {topic_title}",
                "topic_objectives": topic_outline["objectives"],
                "script": "",
            }
            export_pptx_json["slides"].append(topic_objectives_slide)

            for slide in topic_content["slides"]:
                export_pptx_json["slides"].append(slide)

        file_name = file_name.replace(" ", "_")

        file_path = generate_powerpoint_presentation(
            export_pptx_json,
            "DWP Template",
            file_name,
            learning_id,
            extension,
        )

        # Remove temp file and folder after sent successfully back to the user
        # using a background task
        background_tasks.add_task(delete_temp_file, file_path)

        logger.info("Successfully created powerpoint file.")
        # Send powepoint back to UI using FileResponse
        return FileResponse(
            pathlib.Path(file_path),
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            filename=file_name,
        )
    except Exception as error:
        logger.error(
            f"Error occured in write_facilitator_pptx_to_file: {error}"
        )
        raise HTTPException(
            status_code=500,
            detail="An error occurred while create powerpoint document",
        )
